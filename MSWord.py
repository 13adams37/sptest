import re
from copy import deepcopy
import docx
import db

nonBreakSpace = u'\xa0'
serial1_count, serial2_count, main_docs_dict = 0, 0, {}


def save_docx_dict(docs_dict, export_path, object_name, new_name=""):
    for name, doc in dict(docs_dict).items():
        try:
            if new_name:
                name = new_name + " копия"
            doc.save(f"{export_path}"'/'f"{object_name} {name}.docx")
        except PermissionError:
            save_docx_dict(docs_dict, export_path, object_name, name)


def processes_runner(method, args, conn, export_args):
    eval(method)(args)
    try:
        conn.send(main_docs_dict.pop('serials'))
    except KeyError:
        pass
    save_docx_dict(main_docs_dict, *export_args)
    conn.close()


def empty_serial(serial):
    if not serial or serial == "0":
        return "–"
    else:
        return serial


def prepared_text_to_display(element_to_edit):
    element = element_to_edit.copy()
    if element['part'] not in ("", "б/н"):
        element['part'] = f"№{nonBreakSpace}{element['part']}"
    else:
        element['part'] = f"б/н"
    output = f"{element['name']} {element['vendor']} {element['model']} {element['part']}"
    return f"{re.sub(' +', ' ', output)}"


def act_table(elements):
    def fill_new_row_act(main_object, cnt):
        cell = table.add_row().cells
        cell[0].text = str(cnt)
        cell[1].text = main_object['name']
        cell[2].text = main_object['model']
        cell[3].text = main_object['part'] if main_object['part'] else 'б/н'
        cell[4].text = main_object['vendor']
        cell[5].text = main_object['amount']
        cell[6].text = empty_serial(main_object['serial1'])
        cell[7].text = "УФ" if main_object['uv'] else empty_serial(main_object['serial2'])
        cell[8].text = 'CC'
        cell[9].text = '2'

    doc = docx.Document()
    table = doc.add_table(rows=0, cols=10)
    table.style = 'Table Grid'
    counter = 0
    global main_docs_dict

    for item in elements:
        counter += 1
        fill_new_row_act(item, counter)

        if item['table']:
            for item1 in item['table']:
                fill_new_row_act(item1, "")

                if item1['table']:
                    for item2 in item1['table']:
                        fill_new_row_act(item2, "")

    main_docs_dict["АКТ"] = doc


def conclusion_table(elements):
    def fill_conclusion_new_row(main_object, cnt):
        global serial2_count
        cell = table.add_row().cells
        cell[0].text = str(cnt)
        cell[1].text = main_object['name']
        cell[2].text = main_object['model']
        cell[3].text = main_object['part']
        cell[4].text = main_object['vendor']
        cell[5].text = main_object['amount']
        cell[6].text = empty_serial(main_object['serial1'])

        if item['uv']:
            cell[7].text = 'УФ'
        else:
            cell[7].text = empty_serial(main_object['serial2'])

        serial2_count += int(main_object['serial2']) if main_object['serial2'] else 0

    def fill_conclusion_by_row(main_object, cnt, row, this_serial_counter):
        global serial1_count
        cell = table.rows[row].cells
        cell[0].text = str(cnt)
        cell[1].text = main_object['name']
        cell[2].text = main_object['model']
        cell[3].text = main_object['part']
        cell[4].text = main_object['vendor']
        cell[5].text = main_object['amount']
        cell[6].text = empty_serial(main_object['serial1'])
        serial1_count += 1 if main_object['serial1'] else 0

        if item['uv'] and not this_serial_counter:
            cell[7].text = 'УФ'
        else:
            cell[7].text = empty_serial(this_serial_counter)

    doc = docx.Document()
    table = doc.add_table(rows=0, cols=8)
    table.style = 'Table Grid'
    current_row = -1
    global serial1_count, serial2_count, main_docs_dict
    serial1_count = 0
    serial2_count = 0

    for counter_l1, item in enumerate(elements, start=1):
        serial_counter = 0
        table.add_row()
        current_row += 1
        l1_row_to_fill = deepcopy(current_row)

        if item['table']:
            counter_l2 = 0
            counter_l3 = 0
            for item1 in item['table']:
                try:
                    if item1['selected']:
                        counter_l2 += 1
                        l2_serial = 0
                        table.add_row()
                        current_row += 1
                        l2_row_to_fill = deepcopy(current_row)

                        if item1['table']:
                            for sel_item2 in item1['table']:
                                try:
                                    if sel_item2['selected']:
                                        counter_l3 += 1
                                        current_row += 1
                                        fill_conclusion_new_row(sel_item2,
                                                                f"{counter_l1}.{counter_l2}.{counter_l3}")
                                except KeyError:
                                    if sel_item2['serial2'] and not sel_item2['uv']:
                                        l2_serial += int(sel_item2['serial2'])
                                finally:
                                    serial1_count += 1 if sel_item2['serial1'] else 0

                        l2_serial += int(item1['serial2']) if item1['serial2'] != "УФ" and item1['serial2'] else 0
                        fill_conclusion_by_row(item1, f"{counter_l1}.{counter_l2}", l2_row_to_fill, str(l2_serial))
                        serial2_count += l2_serial
                        continue

                except KeyError:
                    if item1['serial2']:
                        serial_counter += int(item1['serial2'])
                    serial1_count += 1 if item1['serial1'] else 0

                if item1['table']:
                    for item2 in item1['table']:
                        try:
                            if item2['selected']:
                                counter_l2 += 1
                                current_row += 1
                                fill_conclusion_new_row(item2,
                                                        f"{counter_l1}.{counter_l2}")

                        except KeyError:
                            if item2['serial2']:
                                serial_counter += int(item2['serial2'])
                        finally:
                            serial1_count += 1 if item2['serial1'] else 0

        serial_counter += int(item['serial2']) if item['serial2'] != "УФ" and item['serial2'] else 0
        fill_conclusion_by_row(item, counter_l1, l1_row_to_fill, str(serial_counter))
        serial2_count += serial_counter

    main_docs_dict.update({"serials": (serial1_count, serial2_count)})
    main_docs_dict['ЗАКЛЮЧЕНИЕ'] = doc


def methods_table(elements):
    def fill_new_row_methods(main_object, cnt, method, method_author):
        cell = table.add_row().cells
        cell[0].text = str(cnt)
        cell[1].text = prepared_text_to_display(main_object)
        cell[2].text = str(method)
        cell[3].text = str(method_author)

    def get_method(name):
        for method_id in iter(db.methods_db):
            method_content = db.methods_db.get(method_id)
            if method_content['type'] == 'По содержанию':
                if name.lower().__contains__(method_content['name'].lower()):
                    return method_content['methods']
            else:
                if method_content['name'].lower() == name:
                    return method_content['methods']
        return ''

    doc = docx.Document()
    table = doc.add_table(rows=0, cols=4)
    table.style = 'Table Grid'
    global main_docs_dict

    for counter_l1, item in enumerate(elements, start=1):
        try:
            author = item['author']
        except KeyError:
            author = ""

        fill_new_row_methods(item, f"{counter_l1}", get_method(item['name'].lower()), author)
        if item['table']:
            for counter_l2, item1 in enumerate(item['table'], start=1):
                fill_new_row_methods(item1, f"{counter_l1}.{counter_l2}", get_method(item1['name'].lower()), author)

                if item1['table']:
                    for counter_l3, item2 in enumerate(item1['table'], start=1):
                        fill_new_row_methods(item2, f"{counter_l1}.{counter_l2}.{counter_l3}",
                                             get_method(item2['name'].lower()), author)

    main_docs_dict['МЕТОДЫ'] = doc


def ims_table(elements):
    def fill_new_row_ims(main_object, cnt):
        cell = table.add_row().cells
        cell[0].text = str(cnt)
        cell[1].text = prepared_text_to_display(main_object)
        cell[3].text = main_object['rgg']
        cell[4].text = main_object['rggpp']

    doc = docx.Document()
    table = doc.add_table(rows=0, cols=5)
    table.style = 'Table Grid'
    global main_docs_dict

    for counter_l1, item in enumerate(elements, start=1):
        fill_new_row_ims(item, f"{counter_l1}")

        if item['table']:
            for counter_l2, item1 in enumerate(item['table'], start=1):
                fill_new_row_ims(item1, f"{counter_l1}.{counter_l2}")

                if item1['table']:
                    for counter_l3, item2 in enumerate(item1['table'], start=1):
                        fill_new_row_ims(item2, f"{counter_l1}.{counter_l2}.{counter_l3}")

    main_docs_dict['СПИСОК ИМС'] = doc
